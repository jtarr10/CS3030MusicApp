'''
This class is for creating a sorted directory from a given directory.
'''

from docx import Document

class Directory:

    #the directory will be sorted by artist by default 
    def __init__(self, songList, sortType='artist'):
        self.songList = [x for x in songList]
        #library will be sorted by the given method
        if(sortType == 'artist'):
            self.sortByArtist()
        elif(sortType == 'title'):
            self.sortByTitle()

    #sorts all songs in the library by artist name
    def sortByArtist(self):
        self.songList.sort(key=lambda x: x.artist)


    #sorts all songs in the library by song title
    def sortByTitle(self):
        self.songList.sort(key=lambda x: x.title)

    #this function will print the directory to a docx file at the given file path
    def print(self, path):
        #we create a document
        doc = Document()
        #put in a fancy heading 
        doc.add_heading('Library Directory:', 0)
        
        #table header setup 
        directory = doc.add_table(rows=1, cols=4)
        directory.style = 'Table Grid'
        header_cells = directory.rows[0].cells
        header_cells[0].text = '#'
        header_cells[1].text = 'Title'
        header_cells[2].text = 'Artist'
        header_cells[3].text = 'Album'

        #table row population from sorted songList
        counter = 1
        for song in self.songList:
            rowCells = directory.add_row().cells
            rowCells[0].text = str(counter)
            rowCells[1].text = song.title
            rowCells[2].text = song.artist
            rowCells[3].text = song.album
            counter += 1

        doc.save(path)
